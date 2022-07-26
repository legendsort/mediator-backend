from docx import *
from docx.shared import *
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml.ns import qn
from django.apps import apps
import datetime
import os
from django.conf import settings
import zipfile
from os.path import basename
from Account.services import GatewayService
from Account.models import RemoteAccount


class SubmissionService:
    def __init__(self, submit):
        self.submit = submit
        self.order = submit.set_order()
        self.DOCUMENT_ROOT_PATH = '/'.join([
            getattr(settings, 'MEDIA_ROOT', '/media'),
            'submissions',
            str(submit.id)
        ])
        self.document = Document()
        section = self.document.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        self.document_name = f"{self.order.id}.docx"
        self.document_path = '/'.join([
            self.DOCUMENT_ROOT_PATH,
            self.document_name
        ])
        self.zip_file_path = f"{self.DOCUMENT_ROOT_PATH}/{str(self.order.id)}.zip"

    def create_information_document(self):
        styles = self.document.styles
        left_style = styles.add_style('left-match', WD_STYLE_TYPE.PARAGRAPH)
        left_style.font.name = u'arial'
        left_style.font.size = Pt(13)
        left_style.paragraph_format.space_after = Pt(1)
        self.document.add_paragraph('Submission information', style='left-match').add_run().add_break()
        self.document.add_paragraph(f"submission title: {self.submit.title}", style='left-match').add_run().add_break()
        self.document.add_paragraph(f"submission article: {self.submit.article.name}", style='left-match').add_run().add_break()
        self.document.add_paragraph(f"submission journal: {self.submit.journal.name}", style='left-match').add_run().add_break()
        self.document.add_paragraph(f"submission publisher: {self.submit.journal.publisher.name}", style='left-match').add_run().add_break()
        self.document.add_paragraph(f"submission abstract: {self.submit.abstract}", style='left-match').add_run().add_break()
        self.document.add_paragraph(f"submission keywords: {self.submit.keywords}", style='left-match').add_run().add_break()
        p = self.document.add_paragraph(f"submission major: {self.submit.major}", style='left-match').add_run()
        p.add_break(WD_BREAK.PAGE)
        self.document.add_paragraph('Account information', style='left-match').add_run().add_break()
        self.document.add_paragraph(f"submission username : {self.submit.user.username}", style='left-match').add_run().add_break()
        # remote account information
        self.create_account_table()
        # resource information
        p = self.document.add_paragraph('', style='left-match').add_run()
        p.add_break()
        p.add_break()
        self.document.add_paragraph('Resource information', style='left-match').add_run().add_break()
        self.create_resource_table()
        # author information
        p = self.document.add_paragraph('', style='left-match').add_run()
        p.add_break()
        p.add_break()
        self.document.add_paragraph('Author information', style='left-match').add_run().add_break()
        self.create_author_table()
        self.save_document()

    def create_account_table(self):
        remote_accounts = RemoteAccount.objects.filter(profile=self.submit.user.profile)
        self.document.styles.add_style('table-type', WD_STYLE_TYPE.TABLE)
        # create account table
        table_header_texts = [
            'No',
            'Email Address',
            'password',
            'Site Address',
            'Journal Name',
            'Available',
        ]
        table = self.document.add_table(1, 6, style='table-type')
        table.style = 'Table Grid'
        table_header = table.rows[0].cells
        index = 0
        for header_text in table_header_texts:
            table_header[index].text = header_text
            table_header[index].paragraphs[0].runs[0].font.bold = True
            table_header[index].paragraphs[0].runs[0].font.size = Pt(12)
            table_header[index].paragraphs[0].runs[0].font.name = u'arial'
            index += 1
        index_number = 1
        for r_account in remote_accounts:
            row_cells = table.add_row().cells
            row_cells[0].text = str(index_number)
            row_cells[1].text = r_account.username
            row_cells[2].text = r_account.password
            row_cells[3].text = r_account.host
            row_cells[4].text = r_account.journal.name if r_account.journal else ''
            row_cells[5].text = 'true' if r_account.is_available else 'false'
            index_number += 1

    def create_author_table(self):
        author_table = self.document.add_table(1, 8, style='table-type')
        author_table_header_texts = [
            'No',
            'Appellation',
            'First Name',
            'Last Name',
            'Email',
            'Reason',
            'Country',
            'Type',
        ]
        author_table_header = author_table.rows[0].cells
        index = 0
        for header_text in author_table_header_texts:
            author_table_header[index].text = header_text
            author_table_header[index].paragraphs[0].runs[0].font.bold = True
            author_table_header[index].paragraphs[0].runs[0].font.size = Pt(12)
            author_table_header[index].paragraphs[0].runs[0].font.name = u'arial'
            index += 1
        index = 1

        for author in self.submit.get_authors():
            row = author_table.add_row().cells
            row[0].text = str(index)
            row[1].text = str(author.appellation)
            row[2].text = str(author.first_name)
            row[3].text = str(author.last_name)
            row[4].text = str(author.email)
            row[5].text = str(author.reason)
            row[6].text = str(author.country.name)
            row[7].text = str(author.type)
            index += 1
        author_table.style = 'Table Grid'

    def create_resource_table(self):
        resource_table = self.document.add_table(1, 3, style='table-type')
        resource_table_header_texts = [
            'No',
            'Requirement',
            'File Name',
        ]
        resource_table_header = resource_table.rows[0].cells
        index = 0
        for header_text in resource_table_header_texts:
            resource_table_header[index].text = header_text
            resource_table_header[index].paragraphs[0].runs[0].font.bold = True
            resource_table_header[index].paragraphs[0].runs[0].font.size = Pt(12)
            resource_table_header[index].paragraphs[0].runs[0].font.name = u'arial'
            index += 1
        index = 1

        for file in self.submit.get_upload_files():
            row = resource_table.add_row().cells
            row[0].text = str(index)
            row[1].text = str(file.requirement.name)
            row[2].text = str(file.name)
            index += 1
        resource_table.style = 'Table Grid'

    def save_document(self):
        try:
            self.create_folder()
            if os.path.exists(self.document_path):
                os.remove(self.document_path)
            self.document.save(self.document_path)
            return True
        except Exception as e:
            print(e)
            return False

    def create_folder(self):
        try:
            if not os.path.isdir(self.DOCUMENT_ROOT_PATH):
                os.mkdir(self.DOCUMENT_ROOT_PATH)
        except Exception as e:
            print('create folder error', e)

    def zip_data(self):
        with zipfile.ZipFile(self.zip_file_path, 'w') as zipObj:
            zipObj.write(self.document_path, basename(self.document_name))
            for file in self.submit.get_upload_files():
                zipObj.write(file.file.path, basename(file.name))
            zipObj.close()
            return True

    def send(self):
        self.make_zip_file()
        try:
            request_zip_file_path = f"{getattr(settings, 'NIS_SEND_DIR_PATH')}/{basename(self.zip_file_path)}"
            if os.path.exists(request_zip_file_path):
                os.remove(request_zip_file_path)
            os.rename(self.zip_file_path, request_zip_file_path)
            gateway_service = GatewayService()
            remote_user_root_path = getattr(settings, 'REMOTE_NEXTCLOUD_USER_ROOT_PATH')
            request_data = {
                'upload_path': f"{remote_user_root_path}/{self.submit.dealer.username}" if self.submit.dealer else '',
                'file_name': basename(self.zip_file_path)
            }
            return gateway_service.send_request(action='upload_resource', data=request_data)
        except Exception as e:
            print('send_error', e)
            return False

    def make_zip_file(self):
        self.create_information_document()
        self.zip_data()
        return self.zip_file_path

    def remove_zip_file(self):
        try:
            if os.path.exists(self.zip_file_path):
                os.remove(self.zip_file_path)
                return True
        except Exception as e:
            print('remove zip file', e)
        return False
