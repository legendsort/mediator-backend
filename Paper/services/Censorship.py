from docx import *
from docx.shared import *
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml.ns import qn
from django.apps import apps
import datetime
import os
from django.conf import settings
import zipfile
from os.path import basename
from Account.services import GatewayService
from Account.models import RemoteAccount


class CensorshipService:
    def __init__(self, order):
        self.order = order
        self.DOCUMENT_ROOT_PATH = '/'.join([
            getattr(settings, 'MEDIA_ROOT', '/media'),
            'order',
            str(order.id)
        ])
        self.document = Document()
        section = self.document.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        self.document_name = f"{self.order.id}.docx"
        self.document_path = '/'.join([
            self.DOCUMENT_ROOT_PATH,
            self.document_name
        ])
        self.zip_file_path = f"{self.DOCUMENT_ROOT_PATH}/{str(self.order.id)}.zip"

    def create_information_document(self):
        styles = self.document.styles
        left_style = styles.add_style('left-match', WD_STYLE_TYPE.PARAGRAPH)
        left_style.font.name = u'arial'
        left_style.font.size = Pt(13)
        left_style.paragraph_format.space_after = Pt(1)
        self.document.add_paragraph('Exchange information', style='left-match').add_run().add_break()
        self.document.add_paragraph(f"Exchange title: {self.exchange.title}", style='left-match').add_run().add_break()
        self.document.add_paragraph(f"Exchange detail: {self.exchange.detail}", style='left-match').add_run().add_break()
        self.document.add_paragraph(f"Exchange purpose: {self.exchange.purpose}", style='left-match').add_run().add_break()
        self.document.add_paragraph(f"Exchange Site: {self.exchange.site_url}", style='left-match').add_run().add_break()
        self.save_document()

    def save_document(self):
        try:
            self.create_folder()
            if os.path.exists(self.document_path):
                os.remove(self.document_path)
            self.document.save(self.document_path)
            return True
        except Exception as e:
            print(e)
            return False

    def create_folder(self):
        try:
            if not os.path.isdir(self.DOCUMENT_ROOT_PATH):
                os.mkdir(self.DOCUMENT_ROOT_PATH)
        except Exception as e:
            print('create folder error', e)

    def zip_data(self):
        with zipfile.ZipFile(self.zip_file_path, 'w') as zipObj:
            zipObj.write(self.document_path, basename(self.document_name))
            file = self.exchange.attachment
            if file:
                print(file.path)
                zipObj.write(file.path, basename(file.name))
                zipObj.close()
            else:
                return True
            return True

    def send(self):
        self.make_zip_file()
        try:
            request_zip_file_path = f"{getattr(settings, 'NIS_SEND_DIR_PATH')}/{basename(self.zip_file_path)}"
            if os.path.exists(request_zip_file_path):
                os.remove(request_zip_file_path)
            os.rename(self.zip_file_path, request_zip_file_path)
            gateway_service = GatewayService()
            remote_user_root_path = getattr(settings, 'REMOTE_NEXTCLOUD_USER_ROOT_PATH')
            request_data = {
                'upload_path': f"{remote_user_root_path}/{self.exchange.dealer.username}" if self.exchange.dealer else '',
                'file_name': basename(self.zip_file_path)
            }
            return gateway_service.send_request(action='upload_resource', data=request_data)
        except Exception as e:
            print('send_error', e)
            return False

    def make_zip_file(self):
        self.create_information_document()
        self.zip_data()
        return self.zip_file_path

    def remove_zip_file(self):
        try:
            if os.path.exists(self.zip_file_path):
                os.remove(self.zip_file_path)
                return True
        except Exception as e:
            print('remove zip file', e)
        return False
